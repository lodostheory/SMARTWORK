"""출장신청서 엑셀 → 새 서식 여비정산서 Word 자동 생성"""
import os, re, datetime
from dataclasses import dataclass, field
from typing import List
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


@dataclass
class TransportEntry:
    date: str = ""
    transport: str = ""
    departure: str = ""
    arrival: str = ""
    grade: str = ""
    amount: str = ""


def get_transport_grade(position: str) -> str:
    return "1" if "교장" in position else "2"


@dataclass
class TravelExpenseData:
    department: str = ""
    position: str = ""
    name: str = ""
    travel_period: str = ""
    destination: str = ""
    basis: str = ""
    purpose: str = ""
    accommodation_limit: str = ""
    accommodation_actual: str = ""
    accommodation_reason: str = ""
    meal_paid: str = ""
    meal_actual: str = ""
    meal_reason: str = ""
    transport: List[TransportEntry] = field(default_factory=list)
    transport_total: str = ""
    year: str = ""
    month: str = ""
    day: str = ""
    applicant: str = ""


FORM_CELLS = {
    "신청일": "W4", "직급": "B10", "성명": "E10", "목적": "I10",
    "시작일": "M10", "시간": "N11", "종료일": "M12",
    "출장지": "U10", "서명": "Y10", "이동사항": "F34",
}


def _v(ws, coord):
    return "" if (v := ws[coord].value) is None else str(v).strip()


def _cell(ws, r, c):
    return "" if (v := ws.cell(r, c).value) is None else str(v).strip()


def parse_period(start, time_r, end):
    wd = ["월", "화", "수", "목", "금", "토", "일"]

    def fmt(d):
        if m := re.match(r"(\d{4})\.(\d{2})\.(\d{2})", d.strip()):
            y, mo, day = m.groups()
            try:
                w = wd[datetime.date(int(y), int(mo), int(day)).weekday()]
            except Exception:
                w = ""
            return f"{y}. {mo}. {day}.({w})", f"{y}-{mo}-{day}"
        return None, None

    s_fmt, s_key = fmt(start)
    e_fmt, e_key = fmt(end)
    ts, te = ([t.strip() for t in re.split(r"[~∼]", time_r)] + ["", ""])[:2]

    if s_fmt and e_fmt:
        if s_key == e_key:
            return f"{s_fmt} {ts} ~ {te}".strip()
        return f"{s_fmt} {ts} ~ {e_fmt} {te}".strip()
    return f"{start} {time_r} ~ {end}".strip()


def load_from_form(xlsx_path, department=""):
    wb = load_workbook(xlsx_path, data_only=True)
    ws = wb.active

    apply_date = _v(ws, FORM_CELLS["신청일"])
    m_date = re.match(r"(\d{4})\.(\d{2})\.(\d{2})", apply_date)
    ay, am, ad = m_date.groups() if m_date else ("", "", "")

    m_pos = _v(ws, FORM_CELLS["직급"])
    m_name = _v(ws, FORM_CELLS["성명"])
    m_purpose = _v(ws, FORM_CELLS["목적"])
    m_dest = _v(ws, FORM_CELLS["출장지"])
    m_sign = _v(ws, FORM_CELLS["서명"])
    m_basis = _v(ws, FORM_CELLS["이동사항"])

    records = []
    for idx in range(1):
        base = 10 + idx * 3
        sd = _cell(ws, base, 13)
        tr = _cell(ws, base + 1, 14)
        ed = _cell(ws, base + 2, 13)
        if not (sd or tr or ed):
            continue

        pos = _cell(ws, base, 2) or m_pos
        name = _cell(ws, base, 5) or m_name
        purpose = _cell(ws, base, 9) or m_purpose
        dest = _cell(ws, base, 21) or m_dest
        sign = _cell(ws, base, 25) or m_sign

        period = parse_period(sd, tr, ed)
        ms = re.match(r"(\d{4})\.(\d{2})\.(\d{2})", sd)
        me = re.match(r"(\d{4})\.(\d{2})\.(\d{2})", ed)

        if ms:
            yy, mm, dd = ms.groups()
            trip_date1 = f"{mm}.{dd}"
        else:
            yy, mm, dd = ay, am, ad
            trip_date1 = f"{mm}.{dd}" if mm and dd else ""

        if me:
            ey, em, ed_val = me.groups()
            trip_date2 = f"{em}.{ed_val}"
            if trip_date1 == trip_date2:
                trip_date2 = trip_date1
        else:
            trip_date2 = trip_date1

        grade = get_transport_grade(pos)
        t_entries = [
            TransportEntry(date=trip_date1, grade=grade),
            TransportEntry(date=trip_date2, grade=grade),
        ]

        today = datetime.date.today()
        ty, tm, td = str(today.year), f"{today.month:02d}", f"{today.day:02d}"

        records.append(TravelExpenseData(
            department=department, position=pos, name=name,
            travel_period=period, destination=dest, basis=m_basis,
            purpose=purpose, transport=t_entries,
            year=ty, month=tm, day=td, applicant=sign or name,
        ))
    return records


# ── Document helpers ──────────────────────────────────────────────────────────

def _border(cell):
    tc = cell._tc.get_or_add_tcPr()
    tcB = tc.find(qn("w:tcBorders"))
    if tcB is None:
        tcB = OxmlElement("w:tcBorders")
        tc.append(tcB)
    for edge in ("top", "left", "bottom", "right"):
        e = tcB.find(qn(f"w:{edge}"))
        if e is None:
            e = OxmlElement(f"w:{edge}")
            tcB.append(e)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:color"), "000000")


def _shade(cell, c):
    tcP = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), c)
    shd.set(qn("w:val"), "clear")
    tcP.append(shd)


def _no_border(cell, top=False, bottom=False, left=False, right=False):
    tc = cell._tc.get_or_add_tcPr()
    tcB = tc.find(qn("w:tcBorders"))
    if tcB is None:
        tcB = OxmlElement("w:tcBorders")
        tc.append(tcB)
    for side in [s for s, f in [("top", top), ("bottom", bottom), ("left", left), ("right", right)] if f]:
        e = tcB.find(qn(f"w:{side}"))
        if e is None:
            e = OxmlElement(f"w:{side}")
            tcB.append(e)
        e.set(qn("w:val"), "none")


def _write(cell, txt, bold=False, al="center", fs=10, sh=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(al, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(txt)
    r.font.name = "맑은 고딕"
    r.font.size = Pt(fs)
    r.bold = bold
    rPr = r._element.get_or_add_rPr()
    rF = OxmlElement("w:rFonts")
    rF.set(qn("w:eastAsia"), "맑은 고딕")
    rPr.append(rF)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if sh:
        _shade(cell, sh)


def _m(tbl, r1, c1, r2, c2):
    return tbl.cell(r1, c1).merge(tbl.cell(r2, c2))


def _set_grid(tbl, ws):
    wt = [int(round(w * 56.6929)) for w in ws]
    tb = tbl._tbl
    tPr = tb.find(qn("w:tblPr"))
    if tPr is None:
        tPr = OxmlElement("w:tblPr")
        tb.insert(0, tPr)
    tL = tPr.find(qn("w:tblLayout"))
    if tL is None:
        tL = OxmlElement("w:tblLayout")
        tPr.append(tL)
    tL.set(qn("w:type"), "fixed")
    tCM = tPr.find(qn("w:tblCellMar"))
    if tCM is None:
        tCM = OxmlElement("w:tblCellMar")
        tPr.append(tCM)
    for s, v in [("left", "60"), ("right", "60"), ("top", "20"), ("bottom", "20")]:
        el = tCM.find(qn(f"w:{s}"))
        if el is None:
            el = OxmlElement(f"w:{s}")
            tCM.append(el)
        el.set(qn("w:w"), v)
        el.set(qn("w:type"), "dxa")
    og = tb.find(qn("w:tblGrid"))
    if og:
        tb.remove(og)
    tG = OxmlElement("w:tblGrid")
    for w in wt:
        gC = OxmlElement("w:gridCol")
        gC.set(qn("w:w"), str(w))
        tG.append(gC)
    tPr.addnext(tG)
    for row in tbl.rows:
        for i, c in enumerate(row.cells):
            if i < len(ws):
                c.width = Mm(ws[i])


# ── Document builder ──────────────────────────────────────────────────────────

def create_doc(data, out):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Mm(20)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("출 장 여 비 정 산 신 청 서")
    r.font.name, r.font.size, r.bold = "맑은 고딕", Pt(20), True
    rPr = r._element.get_or_add_rPr()
    rF = OxmlElement("w:rFonts")
    rF.set(qn("w:eastAsia"), "맑은 고딕")
    rPr.append(rF)

    doc.add_paragraph()

    tbl = doc.add_table(rows=13, cols=10)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    for row in tbl.rows:
        for c in row.cells:
            _border(c)

    cw = [17, 15, 15, 18, 18, 18, 18, 18, 16, 13]
    _set_grid(tbl, cw)

    SHADE = "E7E6E6"

    # Row 0
    _write(tbl.cell(0, 0), "소 속", bold=True, sh=SHADE)
    _m(tbl, 0, 1, 0, 3)
    _write(tbl.cell(0, 1), data.department)
    _write(tbl.cell(0, 4), "직  급\n(직위)", bold=True, sh=SHADE, fs=9)
    _m(tbl, 0, 5, 0, 6)
    _write(tbl.cell(0, 5), data.position)
    _write(tbl.cell(0, 7), "성  명", bold=True, sh=SHADE)
    _m(tbl, 0, 8, 0, 9)
    _write(tbl.cell(0, 8), data.name)

    # Rows 1-3: 출장일정
    _m(tbl, 1, 0, 3, 0)
    _write(tbl.cell(1, 0), "출장\n일정", bold=True, sh=SHADE, fs=9)

    _write(tbl.cell(1, 1), "일   시", bold=True, sh=SHADE, fs=9)
    _m(tbl, 1, 2, 1, 9)
    _write(tbl.cell(1, 2), data.travel_period, al="left")

    _m(tbl, 2, 1, 3, 1)
    _write(tbl.cell(2, 1), "출장지", bold=True, sh=SHADE, fs=9)
    _m(tbl, 2, 2, 3, 4)
    _write(tbl.cell(2, 2), data.destination, al="left")

    _write(tbl.cell(2, 5), "출장근거", bold=True, sh=SHADE, fs=9)
    _m(tbl, 2, 6, 2, 9)
    _write(tbl.cell(2, 6), data.basis, al="left")

    _write(tbl.cell(3, 5), "출장목적", bold=True, sh=SHADE, fs=9)
    _m(tbl, 3, 6, 3, 9)
    _write(tbl.cell(3, 6), data.purpose, al="left")

    # Row 4: 숙박비
    _write(tbl.cell(4, 0), "숙 박 비", bold=True, sh=SHADE)
    _m(tbl, 4, 1, 4, 2)
    _write(tbl.cell(4, 1), "상한액 또는\n지급받은 금액", bold=True, sh=SHADE, fs=9)
    _m(tbl, 4, 3, 4, 4)
    _write(tbl.cell(4, 3), data.accommodation_limit)
    _write(tbl.cell(4, 5), "실제\n소요액", bold=True, sh=SHADE, fs=9)
    _write(tbl.cell(4, 6), data.accommodation_actual)
    _m(tbl, 4, 7, 4, 8)
    _write(tbl.cell(4, 7), "초과지출\n사유", bold=True, sh=SHADE, fs=9)
    _write(tbl.cell(4, 9), data.accommodation_reason, al="left")

    # Row 5: 식비
    _write(tbl.cell(5, 0), "식  비", bold=True, sh=SHADE)
    _m(tbl, 5, 1, 5, 2)
    _write(tbl.cell(5, 1), "지급받은 금액", bold=True, sh=SHADE, fs=9)
    _m(tbl, 5, 3, 5, 4)
    _write(tbl.cell(5, 3), data.meal_paid)
    _write(tbl.cell(5, 5), "실제\n소요액", bold=True, sh=SHADE, fs=9)
    _write(tbl.cell(5, 6), data.meal_actual)
    _m(tbl, 5, 7, 5, 8)
    _write(tbl.cell(5, 7), "초과지출\n사유", bold=True, sh=SHADE, fs=9)
    _write(tbl.cell(5, 9), data.meal_reason, al="left")

    # Rows 6-9: 운임·연료비
    fuel_label = _m(tbl, 6, 0, 9, 0)
    _write(fuel_label, "운 임\n연료비", bold=True, sh=SHADE, fs=9)

    _m(tbl, 6, 1, 6, 2)
    _write(tbl.cell(6, 1), "일 자", bold=True, sh=SHADE)
    _m(tbl, 6, 3, 6, 4)
    _write(tbl.cell(6, 3), "교통편", bold=True, sh=SHADE)
    _write(tbl.cell(6, 5), "출발지", bold=True, sh=SHADE)
    _write(tbl.cell(6, 6), "도착지", bold=True, sh=SHADE)
    _write(tbl.cell(6, 7), "등 급", bold=True, sh=SHADE)
    _m(tbl, 6, 8, 6, 9)
    _write(tbl.cell(6, 8), "금 액", bold=True, sh=SHADE)

    for idx, ri in enumerate([7, 8]):
        if idx < len(data.transport):
            t_entry = data.transport[idx]
            _m(tbl, ri, 1, ri, 2)
            _write(tbl.cell(ri, 1), t_entry.date)
            _m(tbl, ri, 3, ri, 4)
            _write(tbl.cell(ri, 3), t_entry.transport)
            _write(tbl.cell(ri, 5), t_entry.departure)
            _write(tbl.cell(ri, 6), t_entry.arrival)
            _write(tbl.cell(ri, 7), t_entry.grade)
            _m(tbl, ri, 8, ri, 9)
            _write(tbl.cell(ri, 8), t_entry.amount, al="right")
        else:
            _m(tbl, ri, 1, ri, 2)
            _m(tbl, ri, 3, ri, 4)
            _m(tbl, ri, 8, ri, 9)

    # Row 9: 계
    _m(tbl, 9, 1, 9, 2)
    _write(tbl.cell(9, 1), "계", bold=True, sh=SHADE)
    _m(tbl, 9, 3, 9, 7)
    _write(tbl.cell(9, 3), "")
    _m(tbl, 9, 8, 9, 9)
    _write(tbl.cell(9, 8), data.transport_total, al="right", bold=True)

    # Row 10: 안내문
    row10_cell = _m(tbl, 10, 0, 10, 9)
    row10_cell.text = ""
    p1 = row10_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run(
        "「공무원여비규정」제 16 조 제 1 항 및 제 2 항의 규정에 의하여 관계서류를 첨부하여\n"
        "위와같이 여비의 정산을 신청합니다. 끝."
    )
    r1.font.name = "맑은 고딕"
    r1.font.size = Pt(11)
    rPr1 = r1._element.get_or_add_rPr()
    rF1 = OxmlElement("w:rFonts")
    rF1.set(qn("w:eastAsia"), "맑은 고딕")
    rPr1.append(rF1)
    p2 = row10_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run("\n붙임  증빙영수증 1 부.")
    r2.font.name = "맑은 고딕"
    r2.font.size = Pt(11)
    rPr2 = r2._element.get_or_add_rPr()
    rF2 = OxmlElement("w:rFonts")
    rF2.set(qn("w:eastAsia"), "맑은 고딕")
    rPr2.append(rF2)
    _no_border(row10_cell, bottom=True)

    # Row 11: 날짜
    _m(tbl, 11, 0, 11, 9)
    _write(tbl.cell(11, 0), f"{data.year}년   {data.month}월   {data.day}일", al="center", fs=11)
    _no_border(tbl.cell(11, 0), top=True, bottom=True)

    # Row 12: 신청인
    _m(tbl, 12, 0, 12, 9)
    _write(tbl.cell(12, 0), f"신 청 인     성 명     {data.applicant}     (인)", al="center", fs=11, bold=True)
    _no_border(tbl.cell(12, 0), top=True)

    os.makedirs(os.path.dirname(os.path.abspath(out)), exist_ok=True)
    doc.save(out)
    return os.path.abspath(out)


def batch_generate(xlsx_path, output_dir, department=""):
    records = load_from_form(xlsx_path, department=department)
    os.makedirs(output_dir, exist_ok=True)
    saved = []
    for i, data in enumerate(records, start=1):
        suffix = f"{data.name}_{data.year}{data.month}{data.day}" if data.name else f"{i:03d}"
        out_path = os.path.join(output_dir, f"출장여비정산신청서_{suffix}.docx")
        path = create_doc(data, out_path)
        saved.append(path)
    return saved
